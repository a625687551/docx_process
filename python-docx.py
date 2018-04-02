#!/usr/bin/env python3

import docx
import win32com
import os

from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx import Document

from win32com import client

file_path = "C:/Users/rising/Desktop/教育系统/docx/input/math-1.docx"


def read_docx(file_name):
    file = Document(file_path)

    problems = []
    temp = []
    for para in file.paragraphs:
        if para.text.startswith("-----"):
            problems.append(temp)
            temp = []
        else:
            temp.append(para.text)
        print(para.text)
    return problems


def save_docx(problem_list=""):
    document = Document()
    for problem in problem_list:
        for para in problem:
            # 添加文本
            paragraph = document.add_paragraph()
            # 设置字号
            run = paragraph.add_run(para)
            document.add_picture('./media/image001.gif', width=Inches(1.25))  # 添加图片
            run.font.size = Pt(12)

            # 设置中文字体
            run.font.name = u'宋体'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.add_paragraph("---------------------------------", style="Normal")
    # p= document.add_paragraph('A plain paragraph having some ')  # 这里是添加一个段落
    # p.add_run('bold').bold = True  # 这里是在这个段落p里文字some后面添加bold字符
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    document.save("output/demo.docx")


def win_docx(file_name, filenameout="demo.html"):
    word = client.Dispatch("Word.Application")
    # 后台运行，不显示，不警告
    doc = word.Documents.Open(file_name)
    print(os.getcwd())

    # 转换成HTML
    # wc = win32com.client.constants
    # word.ActiveDocument.WebOptions.RelyOnCSS = 1
    # word.ActiveDocument.WebOptions.OptimizeForBrowser = 1
    # word.ActiveDocument.WebOptions.BrowserLevel = 0  # constants.wdBrowserLevelV4
    # word.ActiveDocument.WebOptions.OrganizeInFolder = 0
    # word.ActiveDocument.WebOptions.UseLongFileNames = 1
    # word.ActiveDocument.WebOptions.RelyOnVML = 0
    # word.ActiveDocument.WebOptions.AllowPNG = 1
    # word.ActiveDocument.SaveAs(FileName=filenameout, FileFormat=wc.wdFormatHTML)
    doc.Close()
    word.Quit()


def merge_docx(filename_list, fileout="output/merge.docx"):
    """图片啊"""
    # 初始化
    word = client.Dispatch("Word.Application")
    word.Visible = False
    # 建立一个新的word临时
    newdoc = word.Documents.Add()
    for file in filename_list[::-1]:
        file = os.path.join(os.getcwd(), file)
        newdoc.Application.Selection.Range.InsertFile(file)
        # doc = word.Documents.Open(file)
        # # 获取位置存到末尾
        # range = newdoc.Range()
        # range.InsertAfter(doc.content)
        # doc.Close()

    # 保存文件
    newdoc.SaveAs(os.path.join(os.getcwd(), fileout))

    newdoc.Close()
    word.Quit()


def merge_docx_1(filename_list, fileout="output/merge.docx"):
    """图片啊"""
    target_document = Document()

    for file in filename_list:
        for paragraph in Document(file).paragraphs:
            text = paragraph.text
            target_document.add_paragraph(text)
    # for paragraph in source_document.paragraphs:
    #     text = paragraph.text
    #     target_document.add_paragraph(text)
    target_document.save("output/demo-1.docx")


if __name__ == '__main__':
    # s = read_docx(file_path)
    # save_docx(s)
    # win_docx(file_path)
    merge_docx(
        ["problems/math-1.docx", "problems/math-2.docx", "problems/math-3.docx"])
    # merge_docx_1(
    #     ["C:/Users/rising/Desktop/教育系统/docx/input/math-1.docx", "C:/Users/rising/Desktop/教育系统/docx/input/math-2.docx"])
